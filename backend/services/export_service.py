import os
from ebooklib import epub
from docx import Document
from typing import List, Dict

import re

def strip_html(html_content):
    clean = re.compile('<.*?>')
    return re.sub(clean, '', html_content)

def export_to_epub(novel_title: str, author: str, chapters: List[Dict], output_path: str):
    book = epub.EpubBook()

    # set metadata
    book.set_identifier(f'id_{novel_title}')
    book.set_title(novel_title)
    book.set_language('zh')
    book.add_author(author)

    # create chapters
    epub_chapters = []
    toc = []

    for i, chap_data in enumerate(chapters):
        title = f"Chapter {chap_data.get('chapter_num')}"
        filename = f'chap_{i+1}.xhtml'
        
        c = epub.EpubHtml(title=title, file_name=filename, lang='zh')
        content = chap_data.get('content', '')
        
        # Simple detection if content is HTML
        if '<p>' in content or '<div>' in content or '<br>' in content:
            c.content = f'<h1>{title}</h1>{content}'
        else:
            # Plain text fallback
            content = content.replace('\n', '<br>')
            c.content = f'<h1>{title}</h1><p>{content}</p>'
        
        book.add_item(c)
        epub_chapters.append(c)
        toc.append(c)

    # define Table Of Contents
    book.toc = tuple(toc)

    # add default NCX and Nav file
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())

    # define CSS style
    style = 'BODY { font-family: "Microsoft YaHei", "SimSun", serif; }' # Use fonts more likely to support Chinese
    nav_css = epub.EpubItem(uid="style_nav", file_name="style/nav.css", media_type="text/css", content=style)

    # add CSS file
    book.add_item(nav_css)

    # basic spine
    book.spine = ['nav'] + epub_chapters

    # write to the file
    epub.write_epub(output_path, book, {})
    return output_path

def export_to_docx(novel_title: str, chapters: List[Dict], output_path: str):
    doc = Document()
    doc.add_heading(novel_title, 0)

    for chap in chapters:
        doc.add_heading(f"Chapter {chap.get('chapter_num')}", level=1)
        content = chap.get('content', '')
        # Strip HTML for docx
        if '<' in content and '>' in content:
            content = strip_html(content.replace('<br>', '\n').replace('</p>', '\n'))
        
        doc.add_paragraph(content)

    doc.save(output_path)
    return output_path
