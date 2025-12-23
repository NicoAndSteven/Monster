from fastapi import FastAPI, HTTPException, Body, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles
from starlette.background import BackgroundTask
from pydantic import BaseModel
from .config import settings
from .models.novel import NovelCreate, ChapterGenerate, Asset, PipelineStatus, ChapterUpdate, PlotChoiceRequest, OutlineUpdate, OutlineGenerate
from .utils import storage
from .services import novel_generator, dashboard_service, export_service
from .services.rpa_doubao import doubao_rpa
import os
import json
from docx import Document
from io import BytesIO
import tempfile
import re
import edge_tts
import time

class ImageGenRequest(BaseModel):
    prompt: str

app = FastAPI()

# Mount static files
os.makedirs("static", exist_ok=True)
app.mount("/static", StaticFiles(directory="static"), name="static")

# Mount data storage for serving generated audio/images
os.makedirs(settings.STORAGE_PATH, exist_ok=True)
app.mount("/data", StaticFiles(directory=settings.STORAGE_PATH), name="data")

# CORS configuration
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- Dashboard & Pipeline ---

@app.get("/api/dashboard/stats")
async def get_dashboard_stats():
    # Get file stats
    file_stats = dashboard_service.get_dashboard_stats()
    
    # Load pipeline status from file or use default
    pipeline_data = storage.load_json("pipeline_status.json")
    if not pipeline_data:
        pipeline_data = {
            "status": "processing",
            "current_stage": "image_gen",
            "progress": 65,
            "task_name": "Chapter 5 Generation"
        }
        # Save default to initialize
        storage.save_json("pipeline_status.json", pipeline_data)
    
    # Mock Analytics Data (in a real app, this would come from DB)
    # Use real asset stats from file_stats
    real_asset_types = file_stats.get("asset_types", {})
    
    analytics = {
        "assets_count": {
            "characters": real_asset_types.get("character", 0),
            "scenes": real_asset_types.get("scene", 0),
            "audio_files": real_asset_types.get("audio", 0)
        },
        "distribution": {
            "douyin": [120, 150, 200, 180, 250, 300, 450],
            "xiaohongshu": [80, 90, 110, 105, 130, 140, 160]
        },
        "pipeline": pipeline_data
    }
    
    return {**file_stats, "analytics": analytics}

@app.post("/api/pipeline/update")
async def update_pipeline_status(status: PipelineStatus):
    storage.save_json("pipeline_status.json", status.dict())
    return {"status": "success", "data": status}

@app.post("/api/system/clean_cache")
async def clean_system_cache():
    storage.clear_cache()
    return {"status": "success", "message": "Cache cleaned"}

# --- Background Tasks ---
def generate_and_save_outline(novel_id: str, novel_type: str):
    try:
        outline = novel_generator.generate_outline(novel_type)
        # Load existing to ensure we don't overwrite updates (though rare this early)
        novel_data = storage.load_json(f"novel_{novel_id}.json")
        if novel_data:
            novel_data['outline'] = outline
            storage.save_json(f"novel_{novel_id}.json", novel_data)
            print(f"Outline generated for novel {novel_id}")
    except Exception as e:
        print(f"Failed to generate outline for {novel_id}: {e}")

# --- Novels ---

@app.post("/api/novels")
async def create_novel(novel: NovelCreate, background_tasks: BackgroundTasks):
    novel_dict = novel.dict()
    
    # Initialize empty outline
    novel_dict['outline'] = ""
    
    # Save initial novel data
    storage.save_json(f"novel_{novel.id}.json", novel_dict)
    
    # Schedule Outline Generation if type is provided
    if novel.type:
        background_tasks.add_task(generate_and_save_outline, novel.id, novel.type)
        
    return {"status": "success", "novel": novel_dict}

@app.put("/api/novels/{id}/outline")
async def update_outline(id: str, update: OutlineUpdate):
    novel_data = storage.load_json(f"novel_{id}.json")
    if not novel_data:
        raise HTTPException(status_code=404, detail="Novel not found")
    
    novel_data["outline"] = update.outline
    storage.save_json(f"novel_{id}.json", novel_data)
    return {"status": "success", "outline": update.outline}

@app.post("/api/novels/{id}/outline/generate")
async def regenerate_outline(id: str, background_tasks: BackgroundTasks):
    novel_data = storage.load_json(f"novel_{id}.json")
    if not novel_data:
        raise HTTPException(status_code=404, detail="Novel not found")
    
    novel_type = novel_data.get("type", "General")
    background_tasks.add_task(generate_and_save_outline, id, novel_type)
    return {"status": "success", "message": "Outline generation started"}

@app.get("/api/novels/{id}/relationships")
async def get_relationships(id: str):
    novel_data = storage.load_json(f"novel_{id}.json")
    if not novel_data:
        raise HTTPException(status_code=404, detail="Novel not found")
        
    text_to_analyze = ""
    if novel_data.get("outline"):
        text_to_analyze += f"【大纲】\n{novel_data.get('outline')}\n\n"
        
    # Get latest chapter
    chapter_num = 1
    latest_chapter_content = ""
    while True:
        chap = storage.load_json(f"novel_{id}_chapter_{chapter_num}.json")
        if not chap: break
        if chap.get("content"): latest_chapter_content = chap.get("content")
        chapter_num += 1
        
    if latest_chapter_content:
        text_to_analyze += f"【最新章节内容】\n{latest_chapter_content}"
        
    if not text_to_analyze:
         return {"nodes": [], "links": []}
         
    graph_data = novel_generator.generate_relationship_graph(text_to_analyze)
    return graph_data

@app.get("/api/novels")
async def list_novels():
    files = storage.get_all_files("novel_*.json")
    novels = []
    for f in files:
        if "chapter" not in f: # Exclude chapters
            data = storage.load_json(os.path.basename(f))
            if data and 'title' in data:
                novels.append(data)
    return novels

@app.delete("/api/novels/{id}")
async def delete_novel(id: str):
    # Delete main novel file
    if not storage.delete_file(f"novel_{id}.json"):
        raise HTTPException(status_code=404, detail="Novel not found")
    
    # Delete all chapters
    chapter_files = storage.get_all_files(f"novel_{id}_chapter_*.json")
    for f in chapter_files:
        storage.delete_file(os.path.basename(f))
        
    return {"status": "success", "message": "Novel deleted"}

@app.get("/api/novels/{id}/chapters")
async def list_chapters(id: str):
    files = storage.get_all_files(f"novel_{id}_chapter_*.json")
    chapters = []
    for f in files:
        data = storage.load_json(os.path.basename(f))
        if data:
            chapters.append({
                "id": data.get("chapter_num"),
                "title": f"Chapter {data.get('chapter_num')}", # Simple title for now
                "chapter_num": data.get("chapter_num")
            })
    # Sort by chapter number
    chapters.sort(key=lambda x: x["chapter_num"])
    return chapters

@app.get("/api/novels/{id}/chapters/{chapter_num}")
async def get_chapter(id: str, chapter_num: int):
    data = storage.load_json(f"novel_{id}_chapter_{chapter_num}.json")
    if not data:
        # Return empty template if not exists but novel exists? 
        # Or 404? Let's return empty structure to be safe for frontend
        return {
            "novel_id": id,
            "chapter_num": chapter_num,
            "content": ""
        }
    return data

@app.put("/api/novels/{id}/chapters/{chapter_num}")
async def update_chapter(id: str, chapter_num: int, update: ChapterUpdate):
    filename = f"novel_{id}_chapter_{chapter_num}.json"
    data = storage.load_json(filename)
    if not data:
        data = {
            "novel_id": id,
            "chapter_num": chapter_num,
            "content": "",
            "mode": "manual"
        }
    
    if update.content is not None:
        data["content"] = update.content
    if update.images is not None:
        data["images"] = update.images
        
    storage.save_json(filename, data)
    return {"status": "success", "chapter": data}

from .models.novel import IllustrationGenerate
import math

@app.post("/api/novels/{id}/chapters/{chapter_num}/generate-illustrations")
async def generate_chapter_illustrations(id: str, chapter_num: int):
    # 1. Load Chapter
    filename = f"novel_{id}_chapter_{chapter_num}.json"
    data = storage.load_json(filename)
    if not data or not data.get("content"):
        raise HTTPException(status_code=400, detail="Chapter content is empty")

    content = data.get("content", "")
    # Remove HTML tags for processing
    clean_content = re.sub(r'<[^>]+>', '', content)
    
    # 2. Split content into ~500 char chunks
    # Simple split logic:
    chunk_size = 500
    chunks = []
    current_pos = 0
    while current_pos < len(clean_content):
        end_pos = min(current_pos + chunk_size, len(clean_content))
        # Try to find a sentence break if possible
        if end_pos < len(clean_content):
            # Look for last period in the range
            period_pos = clean_content.rfind('。', current_pos, end_pos)
            if period_pos != -1 and period_pos > current_pos + 300: # Ensure chunk isn't too small
                end_pos = period_pos + 1
        
        chunks.append(clean_content[current_pos:end_pos])
        current_pos = end_pos

    # 3. Generate prompts and images for each chunk
    generated_images = []
    
    # Mock/RPA Image Generation (Since we don't have a real high-speed image gen API)
    # In a real scenario, this would call ComfyUI or Midjourney API
    # Here we use placeholder logic or our existing RPA placeholder
    
    for i, chunk in enumerate(chunks):
        prompt = novel_generator.generate_illustration_prompt(chunk)
        # Using placeholder image service for speed and demo
        # Format: https://placehold.co/512x512?text=Image+for+Part+X
        # In future: call doubao_rpa or dashscope image gen
        
        # Simulating "generated" image URL
        # We'll use a consistent placeholder that looks like a generated image
        # Or better, we can use the prompt as text
        safe_prompt = prompt[:20].replace(" ", "+")
        img_url = f"https://placehold.co/1024x1024/2d3748/cbd5e0?text=Illustration+{i+1}\n{safe_prompt}..."
        
        generated_images.append({
            "id": f"img_{int(tempfile.time.time())}_{i}",
            "prompt": prompt,
            "url": img_url,
            "segment_text": chunk[:50] + "..."
        })
        
    # 4. Save images to chapter
    data["images"] = generated_images
    storage.save_json(filename, data)
    
    return {"status": "success", "images": generated_images}


@app.delete("/api/novels/{id}/chapters/{chapter_num}")
async def delete_chapter(id: str, chapter_num: int):
    filename = f"novel_{id}_chapter_{chapter_num}.json"
    if not storage.delete_file(filename):
        raise HTTPException(status_code=404, detail="Chapter not found")
    return {"status": "success", "message": "Chapter deleted"}

@app.post("/api/novels/{id}/generate")
async def generate_chapter(id: str, chapter: ChapterGenerate):
    # Check if novel exists
    novel_data = storage.load_json(f"novel_{id}.json")
    if not novel_data:
        raise HTTPException(status_code=404, detail="Novel not found")
    
    # 1. Build Context: Assets
    context_parts = []
    
    # Add Outline to context
    if novel_data.get("outline"):
        context_parts.append(f"【小说大纲】\n{novel_data.get('outline')}")

    if chapter.include_assets:
        assets = storage.load_json(f"novel_{id}_assets.json")
        if assets:
            # Filter assets? For now, include all characters and scenes
            # Or maybe just names and roles to save tokens
            relevant_assets = [
                f"{a.get('type').upper()}: {a.get('name')} - {a.get('role') or 'No description'}"
                for a in assets if a.get('type') in ['character', 'scene']
            ]
            if relevant_assets:
                context_parts.append("【相关设定】\n" + "\n".join(relevant_assets))
    
    # 2. Build Context: Previous Chapter
    if chapter.chapter_num > 1:
        prev_chapter_num = chapter.chapter_num - 1
        prev_chapter = storage.load_json(f"novel_{id}_chapter_{prev_chapter_num}.json")
        if prev_chapter and prev_chapter.get("content"):
            content_preview = prev_chapter.get("content")[-chapter.context_window:]
            context_parts.append(f"【前情提要（上一章结尾）】\n...{content_preview}")
    
    # Add Plot Choice
    if chapter.plot_choice:
        context_parts.append(f"【用户选择的剧情走向】\n{chapter.plot_choice}")
            
    full_context = "\n\n".join(context_parts)
    
    # Generate content
    content = novel_generator.generate_chapter_text(
        chapter.prompt or f"Chapter {chapter.chapter_num}",
        mode=chapter.mode,
        context=full_context
    )
    
    # Save chapter
    chapter_data = {
        "novel_id": id,
        "chapter_num": chapter.chapter_num,
        "content": content,
        "mode": chapter.mode
    }
    storage.save_json(f"novel_{id}_chapter_{chapter.chapter_num}.json", chapter_data)
    
    return {"status": "success", "chapter": chapter_data}

@app.post("/api/novels/{id}/plot-choices")
async def get_plot_choices(id: str, request: PlotChoiceRequest):
    print(f"DEBUG: get_plot_choices called with id={id}, request={request}")
    # Check if novel exists
    novel_data = storage.load_json(f"novel_{id}.json")
    if not novel_data:
        raise HTTPException(status_code=404, detail="Novel not found")

    context_parts = []
    
    if novel_data.get("outline"):
        context_parts.append(f"【小说大纲】\n{novel_data.get('outline')}")
    
    # Get previous chapter content for context
    if request.chapter_num > 1:
        prev_chapter_num = request.chapter_num - 1
        prev_chapter = storage.load_json(f"novel_{id}_chapter_{prev_chapter_num}.json")
        if prev_chapter and prev_chapter.get("content"):
             content_preview = prev_chapter.get("content")[-request.context_window:]
             context_parts.append(f"【前情提要】\n...{content_preview}")
    
    full_context = "\n\n".join(context_parts)
    choices = novel_generator.generate_plot_choices(full_context)
    return choices

@app.post("/api/ai/edit")
async def ai_edit_text(request: dict = Body(...)):
    text = request.get("text", "")
    instruction = request.get("instruction", "")
    
    if not text or not instruction:
        raise HTTPException(status_code=400, detail="Missing text or instruction")
        
    result = novel_generator.edit_text(text, instruction)
    if result.startswith("Error:"):
        raise HTTPException(status_code=500, detail=result)
        
    return {"result": result}

@app.post("/api/novels/{id}/assets/{asset_id}/wiki")
async def generate_asset_wiki(id: str, asset_id: int):
    # Load novel and assets
    novel_data = storage.load_json(f"novel_{id}.json")
    assets = storage.load_json(f"novel_{id}_assets.json")
    if not assets:
        raise HTTPException(status_code=404, detail="No assets found")
    
    # Find asset
    target_asset = None
    for asset in assets:
        if asset.get("id") == asset_id:
            target_asset = asset
            break
            
    if not target_asset:
        raise HTTPException(status_code=404, detail="Asset not found")
        
    # Generate Wiki
    wiki_content = novel_generator.generate_wiki_entry(
        name=target_asset.get("name"),
        role=target_asset.get("role", ""),
        basic_info=f"Tags: {', '.join(target_asset.get('tags', []))}",
        context=novel_data.get("outline", "") if novel_data else ""
    )
    
    if wiki_content.startswith("Error:"):
        raise HTTPException(status_code=500, detail=wiki_content)
        
    # Save back
    target_asset["details"] = wiki_content
    storage.save_json(f"novel_{id}_assets.json", assets)
    
    return {"status": "success", "data": target_asset}

@app.get("/api/novels/{id}/export")
async def export_novel(id: str, format: str = "docx"):
    # 1. Load Novel Info
    novel_data = storage.load_json(f"novel_{id}.json")
    if not novel_data:
        raise HTTPException(status_code=404, detail="Novel not found")
        
    # 2. Load All Chapters
    files = storage.get_all_files(f"novel_{id}_chapter_*.json")
    chapters = []
    for f in files:
        data = storage.load_json(os.path.basename(f))
        if data:
            chapters.append(data)
    
    # Sort chapters by number
    chapters.sort(key=lambda x: x.get("chapter_num", 0))

    if format == "epub":
        fd, path = tempfile.mkstemp(suffix=".epub")
        os.close(fd) # Close file descriptor, let ebooklib handle it
        try:
            export_service.export_to_epub(
                novel_title=novel_data.get("title", "Untitled Novel"),
                author="AI Author", # Could add author field later
                chapters=chapters,
                output_path=path
            )
            return FileResponse(
                path,
                filename=f"{novel_data.get('title', 'novel')}.epub",
                media_type="application/epub+zip"
            )
        except Exception as e:
            if os.path.exists(path):
                os.unlink(path)
            raise HTTPException(status_code=500, detail=str(e))

    elif format == "txt":
        content_lines = []
        content_lines.append(novel_data.get("title", "Untitled Novel"))
        content_lines.append(novel_data.get("description", ""))
        content_lines.append("\n" + "="*20 + "\n")
        
        for chapter in chapters:
            chapter_num = chapter.get("chapter_num")
            content = chapter.get("content", "")
            
            # Clean HTML
            clean_content = re.sub(r'</p>', '\n', content)
            clean_content = re.sub(r'<[^>]+>', '', clean_content)
            
            content_lines.append(f"Chapter {chapter_num}")
            content_lines.append(clean_content)
            content_lines.append("\n" + "-"*10 + "\n")
            
        full_text = "\n".join(content_lines)
        
        # Save to temp file
        fd, path = tempfile.mkstemp(suffix=".txt")
        try:
            with os.fdopen(fd, 'w', encoding='utf-8') as tmp:
                tmp.write(full_text)
                
            return FileResponse(
                path,
                filename=f"{novel_data.get('title', 'novel')}.txt",
                media_type="text/plain"
            )
        except Exception as e:
            os.unlink(path)
            raise HTTPException(status_code=500, detail=str(e))

    else: # docx
        # 3. Create Word Document
        doc = Document()
        
        # Title Page
        doc.add_heading(novel_data.get("title", "Untitled Novel"), 0)
        if novel_data.get("description"):
            doc.add_paragraph(novel_data.get("description"))
        doc.add_page_break()
        
        # Chapters
        for chapter in chapters:
            chapter_num = chapter.get("chapter_num")
            content = chapter.get("content", "")
            
            # Clean HTML tags if content is HTML (from Tiptap)
            # Simple regex to strip tags, better to use BeautifulSoup but regex is faster for now
            # Replace <p> with newline, strip others
            clean_content = re.sub(r'</p>', '\n', content)
            clean_content = re.sub(r'<[^>]+>', '', clean_content)
            
            doc.add_heading(f"Chapter {chapter_num}", level=1)
            doc.add_paragraph(clean_content)
            doc.add_page_break()
            
        # Save to temp file
        fd, path = tempfile.mkstemp(suffix=".docx")
        try:
            doc.save(path)
            return FileResponse(
                path,
                filename=f"{novel_data.get('title', 'novel')}.docx",
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        except Exception as e:
            os.unlink(path)
            raise HTTPException(status_code=500, detail=str(e))

# --- Assets (Library) ---

@app.get("/api/novels/{novel_id}/assets")
async def list_assets(novel_id: str):
    assets = storage.load_json(f"novel_{novel_id}_assets.json")
    if not assets:
        assets = [] 
    return assets

@app.post("/api/novels/{novel_id}/assets")
async def create_asset(novel_id: str, asset: Asset):
    filename = f"novel_{novel_id}_assets.json"
    assets = storage.load_json(filename) or []
    
    # Ensure ID is unique or generate one if not provided (though frontend usually provides it)
    # Simple check if exists
    for a in assets:
        if str(a.get("id")) == str(asset.id):
            raise HTTPException(status_code=400, detail="Asset ID already exists")
            
    assets.append(asset.dict())
    storage.save_json(filename, assets)
    return {"status": "success", "asset": asset}

@app.put("/api/novels/{novel_id}/assets/{asset_id}")
async def update_asset(novel_id: str, asset_id: str, asset_update: Asset):
    filename = f"novel_{novel_id}_assets.json"
    assets = storage.load_json(filename) or []
    
    for i, asset in enumerate(assets):
        if str(asset.get("id")) == str(asset_id):
            assets[i] = asset_update.dict()
            storage.save_json(filename, assets)
            return {"status": "success", "asset": asset_update}
            
    raise HTTPException(status_code=404, detail="Asset not found")

@app.delete("/api/novels/{novel_id}/assets/{asset_id}")
async def delete_asset(novel_id: str, asset_id: str):
    filename = f"novel_{novel_id}_assets.json"
    assets = storage.load_json(filename) or []
    
    initial_len = len(assets)
    assets = [a for a in assets if str(a.get("id")) != str(asset_id)]
    
    if len(assets) == initial_len:
        raise HTTPException(status_code=404, detail="Asset not found")
        
    storage.save_json(filename, assets)
    return {"status": "success", "message": "Asset deleted"}

@app.post("/api/generate/image/rpa")
async def generate_image_rpa(request: ImageGenRequest):
    result = await doubao_rpa.generate_image(request.prompt)
    if "error" in result:
        raise HTTPException(status_code=500, detail=result["error"])
    return result

@app.get("/api/media/{chapter_id}/video")
async def get_video(chapter_id: str):
    return {"url": f"https://cdn.example.com/videos/{chapter_id}.mp4"}

@app.post("/api/novels/{id}/analyze-assets")
async def analyze_assets(id: str):
    # 1. Load latest chapter or full text (let's use latest chapter for now for speed)
    # Finding the latest chapter
    chapter_num = 1
    latest_content = ""
    
    while True:
        data = storage.load_json(f"novel_{id}_chapter_{chapter_num}.json")
        if not data:
            break
        if data.get("content"):
            latest_content = data.get("content")
        chapter_num += 1
    
    if not latest_content:
         raise HTTPException(status_code=400, detail="No content to analyze")

    # 2. Load current assets and novel info (for outline)
    assets_filename = f"novel_{id}_assets.json"
    current_assets = storage.load_json(assets_filename) or []
    
    novel_data = storage.load_json(f"novel_{id}.json")
    outline_text = ""
    if novel_data:
        # Concatenate synopsis and outline if available
        parts = []
        if novel_data.get("synopsis"): parts.append(f"简介：{novel_data.get('synopsis')}")
        if novel_data.get("outline"): parts.append(f"大纲：{novel_data.get('outline')}")
        outline_text = "\n".join(parts)

    # 3. Extract updates
    updates = novel_generator.extract_assets_from_text(latest_content, current_assets, outline=outline_text)
    
    # 4. Apply updates
    new_count = 0
    updated_count = 0
    
    for update in updates:
        action = update.get("action", "create")
        name = update.get("name")
        if not name: continue
        
        existing_idx = next((i for i, a in enumerate(current_assets) if a["name"] == name), -1)
        
        new_asset = {
            "id": int(time.time() * 1000) + new_count, # Simple ID generation
            "type": update.get("type", "character"),
            "name": name,
            "role": update.get("role", ""),
            "tags": update.get("tags", []),
            "img": None
        }

        if existing_idx >= 0:
            if action == "update" or action == "create": # Allow create to update if exists
                # Merge logic: preserve ID and Img, update role/tags if provided
                current_assets[existing_idx]["role"] = update.get("role") or current_assets[existing_idx]["role"]
                # Merge tags
                new_tags = set(current_assets[existing_idx].get("tags", []) + update.get("tags", []))
                current_assets[existing_idx]["tags"] = list(new_tags)
                updated_count += 1
        else:
            if action == "create":
                current_assets.append(new_asset)
                new_count += 1

    # 5. Save back
    storage.save_json(assets_filename, current_assets)
    
    return {"status": "success", "new_assets_count": new_count, "updated_assets_count": updated_count}

@app.post("/api/novels/{id}/assets/{asset_name}/refresh")
async def refresh_single_asset(id: str, asset_name: str):
    # 1. Load context (latest chapter or full text)
    chapter_num = 1
    latest_content = ""
    while True:
        data = storage.load_json(f"novel_{id}_chapter_{chapter_num}.json")
        if not data: break
        if data.get("content"): latest_content = data.get("content")
        chapter_num += 1
    
    if not latest_content:
         raise HTTPException(status_code=400, detail="No content to analyze")

    # 2. Load asset
    assets_filename = f"novel_{id}_assets.json"
    current_assets = storage.load_json(assets_filename) or []
    
    asset_idx = next((i for i, a in enumerate(current_assets) if a["name"] == asset_name), -1)
    if asset_idx == -1:
        raise HTTPException(status_code=404, detail="Asset not found")
        
    current_info = current_assets[asset_idx]
    
    # 3. Call AI
    updates = novel_generator.refresh_single_asset(asset_name, latest_content, current_info)
    
    # 4. Update
    if updates:
        current_assets[asset_idx]["role"] = updates.get("role") or current_assets[asset_idx]["role"]
        new_tags = set(current_assets[asset_idx].get("tags", []) + updates.get("tags", []))
        current_assets[asset_idx]["tags"] = list(new_tags)
        
        storage.save_json(assets_filename, current_assets)
        return {"status": "success", "asset": current_assets[asset_idx]}
    else:
        return {"status": "no_change", "asset": current_assets[asset_idx]}

@app.delete("/api/novels/{id}/assets/{asset_id}")
async def delete_asset(id: str, asset_id: int):
    assets_filename = f"novel_{id}_assets.json"
    current_assets = storage.load_json(assets_filename) or []
    
    # Filter out the asset with the given ID
    new_assets = [a for a in current_assets if str(a.get("id")) != str(asset_id)]
    
    if len(new_assets) == len(current_assets):
        raise HTTPException(status_code=404, detail="Asset not found")
    
    storage.save_json(assets_filename, new_assets)
    return {"status": "success", "message": "Asset deleted"}

@app.post("/api/novels/{id}/chapters/{chapter_num}/generate-audio")
async def generate_chapter_audio(id: str, chapter_num: int, request: dict = Body(...)):
    text = request.get("text", "")
    voice = request.get("voice", "zh-CN-XiaoxiaoNeural")
    
    if not text:
         raise HTTPException(status_code=400, detail="Text is empty")
    
    # Directory: data/novels/{id}/audio/
    audio_dir = os.path.join(settings.STORAGE_PATH, f"novel_{id}", "audio")
    os.makedirs(audio_dir, exist_ok=True)
    
    # Filename: {chapter_num}_tts_{timestamp}.mp3
    filename = f"{chapter_num}_tts_{int(time.time())}.mp3"
    filepath = os.path.join(audio_dir, filename)
    
    try:
        # Generate TTS
        communicate = edge_tts.Communicate(text[:5000], voice) # Limit length
        await communicate.save(filepath)
        
        # Add to Assets (Multimedia Warehouse)
        assets_filename = f"novel_{id}_assets.json"
        assets = storage.load_json(assets_filename) or []
        
        # Check if audio asset for this chapter already exists? 
        # Or just append as a new "audio" type asset
        new_asset = {
                    "id": int(time.time() * 1000),
                    "type": "audio",
                    "name": f"第{chapter_num}章配音",
                    "role": f"Chapter {chapter_num} Audio",
                    "tags": ["audio", "tts"],
                    "img": None,
                    "file_path": f"novel_{id}/audio/{filename}" # Relative path for serving
                }
        assets.append(new_asset)
        storage.save_json(assets_filename, assets)
        
        # Serve via static mount? 
        # We need to mount the storage path to serve these files.
        # For now, let's assume we can serve them or just return success.
        # To serve, we might need a specific endpoint or mount.
        # Let's add a generic file server endpoint if needed, or mount 'data'
        
        return {"status": "success", "asset": new_asset}
    except Exception as e:
        print(f"Audio Gen Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/tts")
async def generate_tts(request: dict = Body(...)):
    text = request.get("text", "")
    voice = request.get("voice", "zh-CN-XiaoxiaoNeural")
    
    if not text:
        raise HTTPException(status_code=400, detail="Text is empty")
        
    # Limit text length for safety/speed
    if len(text) > 5000:
        text = text[:5000]

    output_filename = f"tts_{int(time.time())}.mp3"
    output_path = os.path.join("static", output_filename)
    
    try:
        communicate = edge_tts.Communicate(text, voice)
        await communicate.save(output_path)
        return {"url": f"http://localhost:8000/static/{output_filename}"}
    except Exception as e:
        print(f"TTS Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
